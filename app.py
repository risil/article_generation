from langchain_community.llms import CTransformers
from langchain_community.llms import Ollama
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import LLMChain
#from langchain import PromptTemplate
import streamlit as st 
import os
from docx import Document
from docx.shared import Inches
import io
from PIL import Image
import requests

st.title("Article Generator App using Llama 2")
user_input = st.text_input("Please enter the idea/topic for the article you want to generate!")
if st.button("Generate"):
    template = f"""You are a digital marketing and SEO expert and your task is to write article 
	so write an article on the given topic: {user_input}. The article must be under 800 words. 
    The article should be be lengthy."""
    llm_model = Ollama(base_url='http://10.1.1.101:11434', 
                       model="llama3.1:8b"
                       )
    result = llm_model.invoke(template)
    st.info("Your article has been been generated successfully!")
    st.write(result)
    # Prepare the download link
    st.download_button(
    label='Download Word Document',
    data=result,
    file_name='document.docx',
    mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )




def create_word_docx(user_input, paragraph):
    # Create a new Word document
    doc = Document()

    # Add the user input to the document
    doc.add_heading(user_input, level=1)
    doc.add_paragraph(paragraph)

    return doc





